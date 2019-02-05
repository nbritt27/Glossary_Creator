from PyDictionary import PyDictionary
dictionary=PyDictionary()
from docx import Document
from docx.shared import Inches
from tkinter import *
import tkinter
document=Document()

class GlossaryMain:
    def __init__(self, master):
        self.master=master
        master.title("Glossary Creator")
        self.textbox=Entry(master, width=100)
        self.textbox.pack()

        self.button=Button(master, text="Create Glossary", command=self.createGlossary)
        self.button.pack()
    def createGlossary(self):
        inputWords=str(self.textbox.get())
        if(inputWords!="" or None):
            inputWords_split=inputWords.split(", ")
            for i in range(0, len(inputWords_split)):
                print(inputWords_split[i])
                currentWord=dictionary.googlemeaning(str(inputWords_split[i]))
                if currentWord is not None:
                    document.add_paragraph (dictionary.googlemeaning(str(inputWords_split[i])))

                else:
                    secondTrialWord=dictionary.meaning(str(inputWords_split[i]))
                    if secondTrialWord is not None:
                        document.add_paragraph(inputWords_split[i] + " : " + str(list(secondTrialWord)[0]).lower())
                        document.add_paragraph(secondTrialWord[str(list(secondTrialWord)[0])][0])
                    else:
                        document.add_paragraph("Sorry, we can't find this word")
                document.add_paragraph(" ")
            document.save("words.docx")
        else:
            print("The textbox is empty")

root=Tk()
gui=GlossaryMain(root)
root.mainloop()


#inputWords=input("What are the words you want defined? ")
